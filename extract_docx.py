"""Extract text from all .docx files in the dataset and save to UTF-8 text file."""
from docx import Document
import os

dataset_dir = r"c:\Users\kumar\Desktop\Projects\Redrob\dataset\[PUB] India_runs_data_and_ai_challenge\India_runs_data_and_ai_challenge"
output_path = r"c:\Users\kumar\Desktop\Projects\Redrob\docx_content.txt"

with open(output_path, 'w', encoding='utf-8') as out:
    for fname in sorted(os.listdir(dataset_dir)):
        if fname.endswith('.docx'):
            fpath = os.path.join(dataset_dir, fname)
            doc = Document(fpath)
            out.write(f"\n{'='*80}\n")
            out.write(f"FILE: {fname}\n")
            out.write(f"{'='*80}\n")
            for para in doc.paragraphs:
                out.write(para.text + '\n')
            for table in doc.tables:
                for row in table.rows:
                    out.write(" | ".join(cell.text for cell in row.cells) + '\n')
    print("Done! Written to", output_path)
